from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_tutorial_doc():
    doc = Document()
    
    # --- 标题 ---
    title = doc.add_heading('AI 绘画评估脚本使用教程与原理说明', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 第一部分：使用教程 ---
    doc.add_heading('第一部分：使用教程 (How to use)', level=1)
    
    doc.add_heading('步骤 1：准备环境 (安装依赖)', level=2)
    p = doc.add_paragraph('你需要安装一些 Python 库。打开你的终端（Terminal 或 CMD），运行以下命令：')
    code_install = doc.add_paragraph('pip install numpy pandas matplotlib seaborn tqdm opencv-python pillow\npip install mediapipe\npip install torch torchvision\npip install lpips\npip install git+https://github.com/openai/CLIP.git')
    code_install.style = 'No Spacing'
    code_install.runs[0].font.name = 'Courier New'
    code_install.runs[0].font.color.rgb = RGBColor(0, 100, 0) # Dark Green

    doc.add_heading('步骤 2：准备数据文件夹 (关键！)', level=2)
    doc.add_paragraph('脚本的核心依赖是 metadata.csv 文件。你需要按照以下结构整理你的实验文件夹：')
    
    folder_structure = """MyProject/
│
├── scripts/
│   └── evaluate.py          <-- 评估脚本
│
└── outputs/
    └── experiments/
        └── exp_v1_test/     <-- 实验文件夹
            ├── metadata.csv <--- 【必须有这个文件】
            ├── orig_001.jpg <--- 原图
            ├── gen_001.jpg  <--- 生成图
            └── ..."""
    p_struct = doc.add_paragraph(folder_structure)
    p_struct.runs[0].font.name = 'Courier New'

    doc.add_paragraph('metadata.csv 必须包含表头，每一行对应一对“原图-生成图”：')
    
    # 插入 CSV 表格示例
    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'original_image_path'
    hdr_cells[1].text = 'generated_image_path'
    hdr_cells[2].text = 'prompt'
    hdr_cells[3].text = 'seed'
    
    row1 = table.rows[1].cells
    row1[0].text = 'orig_001.jpg'
    row1[1].text = 'gen_001.jpg'
    row1[2].text = 'girl running'
    row1[3].text = '123'

    row2 = table.rows[2].cells
    row2[0].text = 'orig_002.jpg'
    row2[1].text = 'gen_002.jpg'
    row2[2].text = 'boy sitting'
    row2[3].text = '456'

    doc.add_heading('步骤 3：运行脚本', level=2)
    doc.add_paragraph('在终端中，进入 MyProject 根目录，运行以下命令：')
    cmd_run = doc.add_paragraph('python scripts/evaluate.py --exp_dir outputs/experiments/exp_v1_test')
    cmd_run.style = 'No Spacing'
    cmd_run.runs[0].font.name = 'Courier New'
    cmd_run.runs[0].font.bold = True

    doc.add_heading('步骤 4：查看结果', level=2)
    p_res = doc.add_paragraph()
    p_res.add_run('运行结束后，去 ').bold = False
    p_res.add_run('eval_results').bold = True
    p_res.add_run(' 文件夹里看：\n')
    p_res.add_run('1. summary_exp_v1_test.md: 包含通过率、平均分、结论。\n')
    p_res.add_run('2. results_exp_v1_test.csv: 每一张图片的具体打分。\n')
    p_res.add_run('3. charts/: 自动生成的统计图表。')

    # --- 第二部分：原理说明 ---
    doc.add_page_break()
    doc.add_heading('第二部分：原理白话文解释', level=1)

    doc.add_heading('1. 姿态评委 (Pose Consistency)', level=2)
    doc.add_paragraph('负责检查生成的人动作变没变。')
    p_pose = doc.add_paragraph()
    p_pose.add_run('原理：').bold = True
    p_pose.add_run('使用 MediaPipe 技术提取人体 33 个关键点（骨架）。计算原图和生成图对应关节之间的距离误差 (MSE)。\n')
    p_pose.add_run('结论：').bold = True
    p_pose.add_run('MSE 越接近 0，动作还原越准。')

    doc.add_heading('2. 风格评委 (Style/LPIPS)', level=2)
    doc.add_paragraph('负责检查图像的“结构”和“感觉”像不像原图。')
    p_lpips = doc.add_paragraph()
    p_lpips.add_run('原理：').bold = True
    p_lpips.add_run('使用深度神经网络模拟人眼视觉，对比高级特征而不是像素点。\n')
    p_lpips.add_run('结论：').bold = True
    p_lpips.add_run('LPIPS 值越低，视觉结构越相似。')

    doc.add_heading('3. 语义评委 (CLIP Score)', level=2)
    doc.add_paragraph('负责检查生成的图是不是你想要的“意思”（与 Prompt 的匹配度）。')
    p_clip = doc.add_paragraph()
    p_clip.add_run('原理：').bold = True
    p_clip.add_run('使用 CLIP 模型将图片和文字映射到同一个空间，计算它们的余弦相似度。\n')
    p_clip.add_run('结论：').bold = True
    p_clip.add_run('分数越高，说明图片越符合文字描述。')

    # --- 总结表格 ---
    doc.add_heading('总结速查表', level=2)
    table_sum = doc.add_table(rows=5, cols=4)
    table_sum.style = 'Table Grid'
    
    # Header
    h_cells = table_sum.rows[0].cells
    headers = ['指标名称', '中文名', '怎么看好坏？', '实际含义']
    for i, h in enumerate(headers):
        h_cells[i].text = h
        h_cells[i].paragraphs[0].runs[0].font.bold = True

    # Data
    data = [
        ['Pose MSE', '姿态误差', '越小越好 (接近0)', '动作准不准。>0.05 说明变形。'],
        ['LPIPS', '感知差异', '越小越好', '结构像不像。0.3~0.6 是正常范围。'],
        ['CLIP Score', '文本匹配度', '越大越好', '是否符合 Prompt。>0.25 算合格。'],
        ['Style Fidelity', '风格符合度', '越大越好', '含不含特定风格元素。']
    ]

    for i, row_data in enumerate(data):
        cells = table_sum.rows[i+1].cells
        for j, txt in enumerate(row_data):
            cells[j].text = txt

    # --- 常见报错 ---
    doc.add_heading('常见报错与解决', level=1)
    doc.add_paragraph('1. FileNotFoundError: Metadata not found -> 检查路径填错了没。')
    doc.add_paragraph('2. CUDA out of memory -> 显存不足，请在代码中强制使用 CPU。')
    doc.add_paragraph('3. OSError: Can\'t load LPIPS/CLIP -> 网络问题，需检查代理或手动下载权重。')

    # 保存文件
    filename = 'AI绘画评估脚本教程.docx'
    doc.save(filename)
    print(f"文件已生成：{filename}")

if __name__ == "__main__":
    create_tutorial_doc()